"""Generate NOH Maternity Programme — Project Description (Word document)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

FONT = "Arial"
BODY_SIZE = 11
HEADING1_SIZE = 14
HEADING2_SIZE = 12
HEADING3_SIZE = 11
BRAND_TEAL = RGBColor(0x59, 0x95, 0x91)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x60, 0x60, 0x60)

OUT = Path("outputs")
OUT.mkdir(exist_ok=True)
OUTFILE = OUT / "2026-03-02_noh_maternity_programme_project_description.docx"

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

def set_font(run, size=BODY_SIZE, bold=False, italic=False, color=BLACK):
    run.font.name      = FONT
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(4)
    size  = {1: HEADING1_SIZE, 2: HEADING2_SIZE, 3: HEADING3_SIZE}[level]
    color = BRAND_TEAL if level in (1, 2) else BLACK
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    return p

def add_body(doc, text, bold=False, italic=False, space_after=6, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_font(r1, bold=True)
    run = p.add_run(text)
    set_font(run)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D6EDEC")
        cell._tc.get_or_add_tcPr().append(shading)
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = 1
run = p.add_run("NOH Maternity Programme")
set_font(run, size=20, bold=True, color=BRAND_TEAL)

p = doc.add_paragraph()
p.alignment = 1
run = p.add_run("24-Month Cashflow Model — Project Description")
set_font(run, size=14, bold=False, color=BLACK)

p = doc.add_paragraph()
p.alignment = 1
run = p.add_run("Version 1.0  |  2 March 2026  |  NOH Expansion Workbench")
set_font(run, size=10, italic=True, color=GREY)

doc.add_page_break()

# ============================================================
# 1. PURPOSE
# ============================================================
add_heading(doc, "1. Purpose", level=1)
add_body(doc, "This document describes the NOH Maternity Programme cashflow model — what it does, why it was built, the decisions and assumptions it encodes, and how to maintain and update it. It serves as the primary reference for anyone working with the model or the dashboard.")
doc.add_paragraph()
add_body(doc, "The model was built to support commercial negotiations with a hospital group partner, with the objective of demonstrating programme financial viability, quantifying the impact of different partnership terms, and providing a live negotiation tool that can be used in real-time discussions.")

# ============================================================
# 2. CONTEXT
# ============================================================
add_heading(doc, "2. Context", level=1)
add_body(doc, "NOH is establishing a midwifery-led maternity programme at a private hospital in Rustenburg. The programme is designed to:")
add_bullet(doc, "Increase utilisation of existing hospital facilities and personnel")
add_bullet(doc, "Provide a platform for expansion into government and mine health contracts")
add_bullet(doc, "Reduce the caesarean section rate from 42% toward 30% over 24 months")
add_bullet(doc, "Build a sustainable mixed-payer model (medical aid and cash patients)")
doc.add_paragraph()
add_body(doc, "The financial model projects programme cashflow over 24 months from launch, disaggregated by cost category, across five scenarios. A live dashboard makes the model accessible to non-technical partners during negotiations.")

# ============================================================
# 3. WHAT WAS BUILT
# ============================================================
add_heading(doc, "3. What was built", level=1)

add_heading(doc, "3.1 Financial model", level=2)
add_body(doc, "File: src/scenario_model.py")
add_body(doc, "A Python-based disaggregated cashflow model that runs 24 monthly periods for each scenario. Outputs are written to dated CSV files in outputs/.")
doc.add_paragraph()
add_body(doc, "Key model features:", bold=True)
add_bullet(doc, "Linear volume ramp from zero to target births over 24 months")
add_bullet(doc, "CPI price escalation applied monthly to both cash and medical aid prices")
add_bullet(doc, "Disaggregated cost stack: antenatal midwifery, board/facility, MO, anaesthesia, OB pool")
add_bullet(doc, "CS/NVD split with independent ramp toward target CS rate")
add_bullet(doc, "Medical aid share ramp from starting mix to target mix")
add_bullet(doc, "Break-even month calculated per scenario")

add_heading(doc, "3.2 Live dashboard", level=2)
add_body(doc, "File: dashboards/app.py")
add_body(doc, "URL: https://noh-medicare.streamlit.app/")
add_body(doc, "A Streamlit dashboard that loads the latest scenario outputs and includes a Live Negotiation panel with six interactive sliders. The custom scenario recomputes in real time as sliders are adjusted, allowing live what-if analysis during meetings.")

add_heading(doc, "3.3 Code repositories", level=2)
add_table(doc,
    ["Repository", "Location", "Visibility"],
    [
        ["Primary (personal)", "github.com/hpmanyonga/noh-expansion-workbench-", "Public"],
        ["Organisation", "github.com/Network-One/noh-expansion-workbench", "Private"],
    ]
)

add_heading(doc, "3.4 Partner documents", level=2)
add_body(doc, "File: outputs/2026-03-02_noh_dashboard_email_and_guide.docx")
add_body(doc, "A Word document containing the partner email (with dashboard link and negotiation variables table) and the dashboard user guide.")

# ============================================================
# 4. SCENARIOS
# ============================================================
add_heading(doc, "4. Scenarios", level=1)
add_table(doc,
    ["Scenario", "Key assumptions", "Purpose"],
    [
        ["Base", "52 births, 6% CPI, 30% CS target by M12", "Central case for negotiations"],
        ["Conservative", "35 births, 35% CS target, slower MA ramp (18 months)", "Downside volume risk"],
        ["Optimistic", "70 births, 25% CS target, faster MA ramp (10 months)", "Upside case"],
        ["High CS Rate", "CS rate stays at 42% throughout", "Cost of not reducing CS rate"],
        ["Price Stress", "CPI at 4% floor for both prices", "Downside price escalation risk"],
        ["Live Negotiation", "User-defined via dashboard sliders", "Real-time negotiation tool"],
    ]
)

# ============================================================
# 5. FULL ASSUMPTIONS
# ============================================================
add_heading(doc, "5. Model assumptions (Version 1.0, 2 March 2026)", level=1)

add_heading(doc, "5.1 Revenue", level=2)
add_table(doc,
    ["Parameter", "Value", "Notes"],
    [
        ["Cash price per birth (Month 1)", "R 40,000", "Escalates at CPI monthly"],
        ["MA price per birth (Month 1)", "R 50,000", "Escalates at CPI; cannot drop"],
        ["CPI escalation — base", "6% per year", "Monthly compounding"],
        ["CPI escalation — price stress floor", "4% per year", "Minimum allowable"],
        ["NOH management fee", "15% of total receipts", "Open to negotiation"],
        ["Medicare fee", "48.68% of total receipts", "Fixed"],
        ["MA share — start", "20%", ""],
        ["MA share — target", "50%", "Reached by Month 12 (base)"],
    ]
)

add_heading(doc, "5.2 Volume", level=2)
add_table(doc,
    ["Parameter", "Value", "Notes"],
    [
        ["Births at Month 24 (base)", "52/month", ""],
        ["Enrolments at Month 24 (base)", "50/month", ""],
        ["CS rate — start", "42%", "Current hospital rate"],
        ["CS rate — target (base)", "30% by Month 12", "Negotiation objective"],
        ["MO-assisted NVD share", "17.5%", "Midpoint of 15–20% range"],
    ]
)

add_heading(doc, "5.3 Outpatient (antenatal) midwifery", level=2)
add_table(doc,
    ["Item", "Detail", "Cost per episode"],
    [
        ["Long visits", "2 visits × 1.0 hr × R310/hr", "R 620"],
        ["Short visits", "8 visits × 0.5 hr × R310/hr", "R 1,240"],
        ["Total per episode", "10 visits, applied to all births", "R 1,860"],
    ]
)

add_heading(doc, "5.4 In-hospital (intrapartum) midwifery", level=2)
add_body(doc, "Collapsed into the per diem. Midwives are on hospital payroll. The midwifery cost is treated as an allowance embedded in the board rate: R200/day uplift for NVD, R300/day uplift for CS. No separate line item in the model.")
doc.add_paragraph()

add_heading(doc, "5.5 Board / facility (incl. midwifery allowance, medicines, consumables)", level=2)
add_table(doc,
    ["Birth type", "LOS", "Rate/day", "Cost per birth"],
    [
        ["NVD", "24 hrs (1.0 day)", "R 3,200", "R 3,200"],
        ["CS", "48 hrs (2.0 days)", "R 4,900", "R 9,800"],
    ]
)

add_heading(doc, "5.6 MO costs", level=2)
add_table(doc,
    ["Item", "Basis", "Rate"],
    [
        ["Session fee", "1 session per 40 enrolments", "R 2,000/session"],
        ["Birth fee", "MO-assisted NVDs + all CS", "R 3,000/birth"],
    ]
)

add_heading(doc, "5.7 Anaesthesia", level=2)
add_table(doc,
    ["Item", "Basis", "Rate"],
    [
        ["Anaesthesia fee", "All CS births", "R 5,000/CS"],
    ]
)

add_heading(doc, "5.8 OB pool", level=2)
add_body(doc, "Restructured from a retainer + PEPM model to a per-event model reflecting actual OB activity. Rationale: more transparent, easier to audit, and more attractive to younger OBs who prefer predictable income tied to programme activity rather than enrolment counts.")
doc.add_paragraph()
add_table(doc,
    ["Component", "Basis", "Rate", "Rationale"],
    [
        ["Antenatal episode fee", "Per birth (all births)", "R 3,000", "2 OB visits × R1,500 per pregnancy"],
        ["CS delivery fee", "Per CS birth", "R 5,000", "OB is operating surgeon for CS"],
        ["Oversight fee", "Per total delivery per month", "R 750", "Case-mix agnostic; tracks actual workload"],
    ]
)

# ============================================================
# 6. KEY DECISIONS AND RATIONALE
# ============================================================
add_heading(doc, "6. Key decisions and rationale", level=1)
add_table(doc,
    ["Decision", "Rationale"],
    [
        ["Pool model preferred over fee-for-service for OB",
         "FFS (antenatal visits + CS delivery fees alone) exceeds pool cost from Month 6 onwards. Pool is cheaper for the programme and offers income stability attractive to younger OBs. Pool also removes incentive to perform CS."],
        ["Oversight fee tied to total deliveries, not enrolments",
         "Enrolment-based fees expose the programme to falloffs, miscarriages, and referrals. Delivery-based fee reflects actual workload and is case-mix agnostic."],
        ["Intrapartum midwifery collapsed into per diem",
         "Midwives are on hospital payroll. Separating their cost creates administrative complexity without adding analytical value. The per diem uplift (R200 NVD, R300 CS) makes the cost transparent and easy to negotiate."],
        ["CPI escalation applied to both prices",
         "MA price is contractually fixed and escalates at CPI. Cash price follows the same logic. The price_stress scenario uses the 4% CPI floor to model the downside."],
        ["CS rate reduction as a core objective",
         "The high_cs_rate scenario shows a significantly worse cashflow outcome. A shared interest in CS reduction aligns clinical and financial objectives and strengthens the partnership case."],
    ]
)

# ============================================================
# 7. HOW TO UPDATE THE MODEL
# ============================================================
add_heading(doc, "7. How to update the model", level=1)

add_heading(doc, "7.1 Changing assumptions", level=2)
add_body(doc, "All assumptions are defined in the BASE dictionary in src/scenario_model.py. To change an assumption:")
add_bullet(doc, "Open src/scenario_model.py")
add_bullet(doc, "Find the parameter in the BASE dictionary")
add_bullet(doc, "Update the value")
add_bullet(doc, "Run: python src/scenario_model.py")
add_bullet(doc, "New output CSVs are written to outputs/ with today's date")
add_bullet(doc, "The dashboard picks up the latest files automatically on next load")
doc.add_paragraph()

add_heading(doc, "7.2 Adding a new scenario", level=2)
add_body(doc, "Add an entry to the SCENARIOS dictionary in src/scenario_model.py using the BASE dictionary as a base and overriding only the parameters that differ. Then add the scenario name and colour to SCENARIO_COLORS and SCENARIO_LABELS in dashboards/app.py.")
doc.add_paragraph()

add_heading(doc, "7.3 Deploying updates", level=2)
add_body(doc, "After making changes:")
add_bullet(doc, "Run: python src/scenario_model.py  (regenerates CSVs)")
add_bullet(doc, "Commit: git add . && git commit -m 'description of change'")
add_bullet(doc, "Push: ~/Desktop/push.sh  (paste PAT when prompted)")
add_bullet(doc, "Streamlit redeploys automatically within 1–2 minutes")
doc.add_paragraph()

add_heading(doc, "7.4 File structure", level=2)
add_table(doc,
    ["File / folder", "Purpose"],
    [
        ["src/scenario_model.py", "All model logic and scenario definitions"],
        ["dashboards/app.py", "Streamlit dashboard"],
        ["outputs/YYYY-MM-DD_scenario_all.csv", "Combined scenario data (all scenarios, all months)"],
        ["outputs/YYYY-MM-DD_scenario_summary.csv", "Month 12 and Month 24 summary per scenario"],
        ["outputs/YYYY-MM-DD_scenario_<name>.csv", "Individual scenario files"],
        ["docs/", "Project documentation and scripts"],
        ["requirements.txt", "Python dependencies for deployment"],
    ]
)

# ============================================================
# 8. NEXT STEPS
# ============================================================
add_heading(doc, "8. Suggested next steps", level=1)
add_table(doc,
    ["Priority", "Action"],
    [
        ["High", "Review all assumptions with hospital group at first negotiation meeting"],
        ["High", "Update model with agreed terms after negotiation and rerun scenarios"],
        ["Medium", "Add break-even callout prominently to dashboard"],
        ["Medium", "Add revenue waterfall chart (gross receipts to net cashflow)"],
        ["Medium", "Model government and mine contract scenarios as volumes are confirmed"],
        ["Low", "Add cost per birth metric to dashboard for unit economics discussion"],
        ["Low", "Add download button for scenario data export"],
    ]
)

doc.save(str(OUTFILE))
print(f"Saved: {OUTFILE}")
